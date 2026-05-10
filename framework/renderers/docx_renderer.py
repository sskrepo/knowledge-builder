"""DOCX renderer — produces a real Word document via python-docx."""
from __future__ import annotations
import io
import logging
from pathlib import Path
from ._base import BaseRenderer

log = logging.getLogger(__name__)


class DocxRenderer(BaseRenderer):
    name = "docx"
    output_extension = ".docx"

    def render(self, data: dict, template: str | Path | None = None) -> bytes:
        try:
            from docx import Document
        except ImportError:
            log.warning("python-docx not installed; producing placeholder bytes")
            return self._stub_bytes(data)

        if template and Path(template).exists():
            doc = Document(str(template))
        else:
            doc = Document()

        # Title
        if "title" in data:
            doc.add_heading(data["title"], level=0)
        if "subtitle" in data:
            doc.add_paragraph(data["subtitle"])

        # Sections
        for section_name, body in data.get("sections", {}).items():
            doc.add_heading(section_name, level=1)
            if isinstance(body, list):
                for item in body:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(body, dict):
                for k, v in body.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k}: ").bold = True
                    p.add_run(str(v))
            else:
                doc.add_paragraph(str(body))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _stub_bytes(data: dict) -> bytes:
        import json
        return ("STUB DOCX (install python-docx to render real)\n" +
                json.dumps(data, indent=2, default=str)).encode("utf-8")
